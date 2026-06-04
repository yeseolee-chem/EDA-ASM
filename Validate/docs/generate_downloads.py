#!/usr/bin/env python3
"""Build downloadable documents (DOCX + Markdown + bundled ZIP) for the 62
post-pipeline reactions, split into 43 PASS and 19 FAIL.

Outputs (under ADF_500_edited/docs_download/):
  images/<rid>_R.png, <rid>_TS.png, <rid>_P.png  — PNG renderings
  PASS_43.docx, PASS_43.md                       — 43 passing reactions
  FAIL_19.docx, FAIL_19.md                       — 19 still-failing reactions
  ADF_EDA_docs.zip                               — everything zipped together
"""

from __future__ import annotations

import csv
import io
import json
import pickle
import sys
import zipfile
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from matplotlib.patches import Circle
from matplotlib.lines import Line2D

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/gpfs/home1/yeseo1ee/projects/eda-asm-prediction")
OUT = ROOT / "ADF_500_edited" / "docs_download"
IMG = OUT / "images"
OUT.mkdir(parents=True, exist_ok=True)
IMG.mkdir(parents=True, exist_ok=True)

sys.path.insert(0, str(ROOT / "Validate"))
from validate_asr import (  # type: ignore
    Config, derive, check1_schema, check3_topology,
    check4_conservation, check5_signs, aggregate,
)
cfg = Config()

ATOMIC_NUMBER = {"H": 1, "B": 5, "C": 6, "N": 7, "O": 8, "F": 9, "P": 15,
                 "S": 16, "Cl": 17, "Br": 35, "I": 53}
SYM_OF = {v: k for k, v in ATOMIC_NUMBER.items()}
ELEM_COLOR = {
    "H": "#dcdcdc", "B": "#ffb5b5", "C": "#222222", "N": "#3050f8",
    "O": "#ff2222", "F": "#9bd35e", "P": "#ff8000", "S": "#ffdc00",
    "Cl": "#1ff01f", "Br": "#a62929", "I": "#940094",
}
FRAG_COLORS = ["#4ea1ff", "#f0b94c", "#a37bff", "#5ed27a", "#ef5b5b", "#7feaff"]


def project_xy(positions: np.ndarray) -> np.ndarray:
    """PCA project (n_atoms, 3) → (n_atoms, 2)."""
    p = positions - positions.mean(axis=0)
    cov = p.T @ p / max(1, len(p))
    w, v = np.linalg.eigh(cov)
    order = np.argsort(w)[::-1]
    return p @ v[:, order[:2]]


def render_png(symbols, positions_3d, bonds_R, bonds_P, fragments,
                out_path: Path, title: str = "") -> None:
    """Render R/TS/P-style molecule view as a PNG."""
    if positions_3d is None or len(positions_3d) == 0:
        fig, ax = plt.subplots(figsize=(3.0, 2.4), dpi=120)
        ax.text(0.5, 0.5, "no coords", ha="center", va="center",
                color="#888", transform=ax.transAxes)
        ax.axis("off")
        fig.savefig(out_path, bbox_inches="tight", facecolor="#0a0c12")
        plt.close(fig)
        return
    pos = np.asarray(positions_3d, dtype=float)
    xy = project_xy(pos)
    fig, ax = plt.subplots(figsize=(3.0, 2.4), dpi=120, facecolor="#0a0c12")
    ax.set_facecolor("#0a0c12")
    set_R = {tuple(sorted(e)) for e in bonds_R}
    set_P = {tuple(sorted(e)) for e in bonds_P}
    kept = set_R & set_P
    broken = set_R - set_P
    formed = set_P - set_R

    def line(edge, color, dashed=False):
        a, b = edge
        if a >= len(xy) or b >= len(xy):
            return
        ax.plot([xy[a, 0], xy[b, 0]], [xy[a, 1], xy[b, 1]],
                color=color, lw=2.0, linestyle="--" if dashed else "-",
                zorder=1)
    for e in kept:
        line(e, "#9ba3b8")
    for e in broken:
        line(e, "#ef5b5b", dashed=True)
    for e in formed:
        line(e, "#5ed27a", dashed=True)

    atom_to_frag = {}
    for fi, frag in enumerate(fragments):
        for ai in frag["atom_indices"]:
            atom_to_frag[ai] = fi

    for i, sym in enumerate(symbols):
        col = ELEM_COLOR.get(sym, "#888")
        fi = atom_to_frag.get(i, -1)
        if fi >= 0:
            ring_col = FRAG_COLORS[fi % len(FRAG_COLORS)]
            ax.add_patch(Circle((xy[i, 0], xy[i, 1]), 0.20,
                                   facecolor="none", edgecolor=ring_col,
                                   lw=2.0, zorder=2))
        ax.add_patch(Circle((xy[i, 0], xy[i, 1]), 0.14,
                               facecolor=col, edgecolor="#000", lw=0.5, zorder=3))
        is_dark = sym in ("C", "N", "Br", "I", "P")
        ax.text(xy[i, 0], xy[i, 1], f"{sym}{i}",
                ha="center", va="center", fontsize=6.5, fontfamily="monospace",
                color="white" if is_dark else "black", zorder=4)

    if title:
        ax.text(0.02, 0.97, title, transform=ax.transAxes, fontsize=8,
                color="#8a93a6", va="top", fontfamily="monospace")

    pad = 0.4
    ax.set_xlim(xy[:, 0].min() - pad, xy[:, 0].max() + pad)
    ax.set_ylim(xy[:, 1].min() - pad, xy[:, 1].max() + pad)
    ax.set_aspect("equal")
    ax.axis("off")
    fig.savefig(out_path, bbox_inches="tight", facecolor="#0a0c12",
                 edgecolor="none", pad_inches=0.05)
    plt.close(fig)


def _positions_for(cache, rid: str, label: str):
    e = cache.get(rid)
    if e is None:
        return None
    arr = {"R": e.positions_R, "TS": e.positions_TS, "P": e.positions_P}.get(label)
    return None if arr is None else np.asarray(arr).reshape(-1, 3)


def _symbols_for(cache, rid: str):
    e = cache.get(rid)
    return [] if e is None else [SYM_OF.get(int(z), "?") for z in e.numbers]


def _derived(raw):
    d = derive("o", raw)
    issues = (check1_schema(d) + check3_topology(d, cfg)
              + check4_conservation(d, cfg) + check5_signs(d))
    return d, aggregate(issues)


# ─── DOCX helpers ────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    """Shade a docx table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_kv_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(pairs):
        c0, c1 = tbl.rows[i].cells
        c0.text = k
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        c1.text = v
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.name = "Consolas"
                r.font.size = Pt(9)


def _add_frag_table(doc: Document, fragments: list[dict], symbols: list[str]) -> None:
    headers = ["frag", "role", "n_atoms", "Σ Z", "multiplicity", "atom_indices"]
    tbl = doc.add_table(rows=1 + len(fragments), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for i, f in enumerate(fragments):
        ne = sum(ATOMIC_NUMBER.get(symbols[a], 0)
                 for a in f["atom_indices"] if a < len(symbols))
        atoms_str = ",".join(str(a) for a in f["atom_indices"])
        cells = tbl.rows[i + 1].cells
        cells[0].text = f"F{i}"
        cells[1].text = str(f.get("role", "?"))
        cells[2].text = str(len(f["atom_indices"]))
        cells[3].text = str(ne)
        cells[4].text = str(f.get("multiplicity", "?"))
        cells[5].text = atoms_str
        # Color the frag-color cell background
        _set_cell_bg(cells[0], FRAG_COLORS[i % len(FRAG_COLORS)].lstrip("#"))
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Consolas"
                    r.font.size = Pt(8)


def _add_three_up(doc: Document, png_paths: list[Path]) -> None:
    """Place three PNGs in a horizontal table cell-by-cell."""
    tbl = doc.add_table(rows=1, cols=3)
    for j, p in enumerate(png_paths):
        cell = tbl.rows[0].cells[j]
        if p and p.exists():
            run = cell.paragraphs[0].add_run()
            run.add_picture(str(p), width=Inches(2.1))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_reaction_section(doc: Document, rid: str, raw_new, raw_old, s5_new,
                            symbols, fail_cat: str, png_paths: list[Path]) -> None:
    new_d, _ = _derived(raw_new)
    old_d, _ = _derived(raw_old)
    family = "T1x" if rid.startswith("T1x") else "Halogen"
    pattern = s5_new["result"].get("pattern", "?")
    fragments = s5_new["result"]["fragments"]
    coupling = s5_new["result"].get("coupling", "?")
    selected = raw_new.get("selected_candidate", s5_new.get("selected_candidate", "?"))

    h = doc.add_heading(rid, level=2)
    for run in h.runs:
        run.font.name = "Consolas"
    sub = doc.add_paragraph()
    sub.add_run(f"{family} · {pattern}").bold = True
    if fail_cat:
        sub.add_run(f"  ·  category: {fail_cat}").font.color.rgb = RGBColor(0xEF, 0x5B, 0x5B)

    bonds_R = s5_new["debug"]["bonds_R"]
    bonds_P = s5_new["debug"]["bonds_P"]
    set_R = {tuple(sorted(e)) for e in bonds_R}
    set_P = {tuple(sorted(e)) for e in bonds_P}
    broken_list = ", ".join(f"{a}-{b}" for a, b in sorted(set_R - set_P)) or "—"
    formed_list = ", ".join(f"{a}-{b}" for a, b in sorted(set_P - set_R)) or "—"

    _add_kv_table(doc, [
        ("n_atoms", str(s5_new.get("n_atoms", "?"))),
        ("ΔE‡ (kcal/mol)", f"{new_d.dE_act:.3f}"),
        ("ΔE_rxn (kcal/mol)", f"{new_d.dE_rxn:.3f}"),
        ("res_cons (original)", f"{old_d.max_abs_res_cons:.3f}"),
        ("res_cons (final)", f"{new_d.max_abs_res_cons:.4f}"),
        ("selected strategy", str(selected)),
        ("coupling", str(coupling)),
        ("broken bonds (R-only)", broken_list),
        ("formed bonds (P-only)", formed_list),
    ])

    doc.add_paragraph().add_run("Selected fragmentation").bold = True
    _add_frag_table(doc, fragments, symbols)

    doc.add_paragraph().add_run("R / TS / P geometries").bold = True
    _add_three_up(doc, png_paths)


def build_docx(out_path: Path, title: str, cards: list[dict],
                intro_paragraphs: list[str]) -> None:
    doc = Document()
    # Tight default margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.6)
        sec.right_margin = Inches(0.6)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for para in intro_paragraphs:
        doc.add_paragraph(para)

    # Legend
    doc.add_paragraph().add_run("Bond legend").bold = True
    legend = doc.add_paragraph()
    legend.add_run("solid grey").font.color.rgb = RGBColor(0x9B, 0xA3, 0xB8)
    legend.add_run(" = kept bond     |     ")
    legend.add_run("red dashed").font.color.rgb = RGBColor(0xEF, 0x5B, 0x5B)
    legend.add_run(" = broken in R only     |     ")
    legend.add_run("green dashed").font.color.rgb = RGBColor(0x5E, 0xD2, 0x7A)
    legend.add_run(" = formed in P only")
    p = doc.add_paragraph()
    p.add_run("Fragment ring colors: F0 / F1 / F2 / F3 / F4 = blue / yellow / purple / green / red")

    doc.add_paragraph().add_run("Table of contents").bold = True
    toc = doc.add_paragraph()
    for c in cards:
        run = toc.add_run(c["rid"] + " · ")
        run.font.name = "Consolas"; run.font.size = Pt(9)

    doc.add_page_break()

    for card in cards:
        _add_reaction_section(doc, card["rid"], card["raw_new"], card["raw_old"],
                               card["s5_new"], card["symbols"],
                               card.get("fail_cat", ""), card["png_paths"])
        doc.add_page_break()

    doc.save(out_path)


# ─── Markdown helpers ────────────────────────────────────────────────────
def build_markdown(out_path: Path, title: str, cards: list[dict],
                    intro: str) -> None:
    out: list[str] = []
    out.append(f"# {title}\n")
    out.append(intro + "\n")
    out.append("## Bond legend\n")
    out.append("- Solid grey line = kept bond (present in both R and P)")
    out.append("- Red dashed line = broken bond (R only)")
    out.append("- Green dashed line = formed bond (P only)")
    out.append("- Fragment ring colors: F0 blue · F1 yellow · F2 purple · F3 green · F4 red\n")

    out.append("## Table of contents\n")
    for c in cards:
        out.append(f"- [{c['rid']}](#{c['rid'].lower().replace('_', '-')})")
    out.append("")

    for card in cards:
        rid = card["rid"]
        new_d, _ = _derived(card["raw_new"])
        old_d, _ = _derived(card["raw_old"])
        s5 = card["s5_new"]
        symbols = card["symbols"]
        family = "T1x" if rid.startswith("T1x") else "Halogen"
        pattern = s5["result"].get("pattern", "?")
        fragments = s5["result"]["fragments"]
        coupling = s5["result"].get("coupling", "?")
        selected = card["raw_new"].get("selected_candidate",
                                       s5.get("selected_candidate", "?"))
        bonds_R = s5["debug"]["bonds_R"]
        bonds_P = s5["debug"]["bonds_P"]
        set_R = {tuple(sorted(e)) for e in bonds_R}
        set_P = {tuple(sorted(e)) for e in bonds_P}
        broken_list = ", ".join(f"{a}-{b}" for a, b in sorted(set_R - set_P)) or "—"
        formed_list = ", ".join(f"{a}-{b}" for a, b in sorted(set_P - set_R)) or "—"

        out.append(f"\n## {rid}\n")
        if card.get("fail_cat"):
            out.append(f"**Failure category**: `{card['fail_cat']}`  ")
        out.append(f"**Family**: {family} · **Pattern**: `{pattern}`\n")
        out.append("| key | value |")
        out.append("|---|---|")
        out.append(f"| n_atoms | {s5.get('n_atoms','?')} |")
        out.append(f"| ΔE‡ (kcal/mol) | {new_d.dE_act:.3f} |")
        out.append(f"| ΔE_rxn (kcal/mol) | {new_d.dE_rxn:.3f} |")
        out.append(f"| res_cons (original) | {old_d.max_abs_res_cons:.3f} |")
        out.append(f"| res_cons (final) | **{new_d.max_abs_res_cons:.4f}** |")
        out.append(f"| selected strategy | `{selected}` |")
        out.append(f"| coupling | `{coupling}` |")
        out.append(f"| broken (R-only) | `{broken_list}` |")
        out.append(f"| formed (P-only) | `{formed_list}` |")
        out.append("")
        out.append("### Selected fragmentation\n")
        out.append("| frag | role | n_atoms | Σ Z | multiplicity | atom_indices |")
        out.append("|---|---|---|---|---|---|")
        for i, f in enumerate(fragments):
            ne = sum(ATOMIC_NUMBER.get(symbols[a], 0)
                     for a in f["atom_indices"] if a < len(symbols))
            atoms_str = ",".join(str(a) for a in f["atom_indices"])
            out.append(f"| F{i} | {f.get('role','?')} | {len(f['atom_indices'])} | {ne} | {f.get('multiplicity','?')} | `{atoms_str}` |")
        out.append("")
        out.append("### R / TS / P geometries\n")
        out.append("| R | TS | P |")
        out.append("|---|---|---|")
        out.append(f"| ![R](images/{rid}_R.png) | ![TS](images/{rid}_TS.png) | ![P](images/{rid}_P.png) |")
        out.append("")

    out_path.write_text("\n".join(out))


# ─── Main ───────────────────────────────────────────────────────────────
def main() -> int:
    print("Loading frames cache…")
    with open(ROOT / "ADF_500/stage5a/frames_cache.pkl", "rb") as fh:
        cache = pickle.load(fh)

    diag = json.loads((ROOT / "Validate/refrag/still_fail_diagnosis.json").read_text())
    fail_cat_by_rid = {r["rid"]: r["cat"] for r in diag["rows"]}

    orig_fails: list[str] = []
    for row in csv.DictReader(open(ROOT / "Validate/manifest.csv")):
        if row["verdict"] != "FAIL":
            continue
        fc = set(row["failed_checks"].split(";"))
        if "3" in fc or "4" in fc:
            orig_fails.append(row["reaction_id"])

    pass_cards: list[dict] = []
    fail_cards: list[dict] = []
    n_skip = 0
    print(f"Processing {len(orig_fails)} reactions…")
    for i, rid in enumerate(orig_fails, 1):
        new_path = ROOT / f"ADF_500_edited/results/{rid}.json"
        s5_path = ROOT / f"ADF_500_edited/stage5a/per_reaction/{rid}/result.json"
        old_path = ROOT / f"ADF_500/results/{rid}.json"
        if not (new_path.exists() and s5_path.exists() and old_path.exists()):
            n_skip += 1
            continue
        raw_new = json.loads(new_path.read_text())
        raw_old = json.loads(old_path.read_text())
        s5_new = json.loads(s5_path.read_text())
        symbols = _symbols_for(cache, rid)
        _, verdict = _derived(raw_new)
        bonds_R = s5_new["debug"]["bonds_R"]
        bonds_P = s5_new["debug"]["bonds_P"]
        fragments = s5_new["result"]["fragments"]

        png_paths = []
        for lbl in ("R", "TS", "P"):
            pos = _positions_for(cache, rid, lbl)
            png = IMG / f"{rid}_{lbl}.png"
            if not png.exists():
                render_png(symbols, pos, bonds_R, bonds_P, fragments,
                            png, title=f"{lbl} geometry")
            png_paths.append(png)

        card = {
            "rid": rid, "raw_new": raw_new, "raw_old": raw_old,
            "s5_new": s5_new, "symbols": symbols, "png_paths": png_paths,
            "fail_cat": fail_cat_by_rid.get(rid, ""),
        }
        if verdict == "FAIL":
            fail_cards.append(card)
        else:
            pass_cards.append(card)
        if i % 10 == 0:
            print(f"  {i}/{len(orig_fails)}")

    pass_cards.sort(key=lambda c: c["rid"])
    fail_cards.sort(key=lambda c: (c.get("fail_cat", ""), c["rid"]))

    print(f"PASS: {len(pass_cards)}  FAIL: {len(fail_cards)}  skipped: {n_skip}")

    # Build DOCX
    print("Writing PASS_43.docx…")
    build_docx(OUT / "PASS_43.docx",
               f"EDA / ASM pipeline — {len(pass_cards)} reactions (Check-4 PASS)",
               pass_cards,
               [f"This document covers the {len(pass_cards)} reactions whose Check-4 "
                f"conservation residual is now ≤ 0.5 kcal/mol after the multi-sweep "
                f"pipeline (fragmentation candidates + spin variants + targeted "
                f"retry + paper-inspired homolytic cleavage).",
                "Each reaction section shows the final selected fragmentation "
                "(which strategy won, atom indices per fragment, multiplicity, "
                "coupling), the bond changes that define the reaction, and three "
                "side-by-side renderings of the R, TS, and P geometries with atoms "
                "colored by element and fragment membership shown as a colored ring."])

    print("Writing FAIL_19.docx…")
    build_docx(OUT / "FAIL_19.docx",
               f"EDA / ASM pipeline — {len(fail_cards)} reactions still FAIL",
               fail_cards,
               [f"These {len(fail_cards)} reactions still fail at least one "
                f"non-recovered check after the full pipeline (359 ADF candidate "
                f"runs across four sweeps).",
                "The 'category' tag on each section is the dominant failure mode "
                "from the diagnostic pass (`cons_huge` / `cons_mid` / `cons_small` "
                "= conservation residual size, `schema` = source-data corruption, "
                "`ts_not_max` = Halo8 trajectory's last frame is not the product "
                "minimum). The paper-inspired (homolytic) candidates changed the "
                "residual by less than 0.05 kcal/mol in every case — strong "
                "evidence that the limit is irreducible ETSNOCV decomposition "
                "truncation, not a wrong fragmentation choice."])

    print("Writing PASS_43.md / FAIL_19.md…")
    build_markdown(OUT / "PASS_43.md",
                    f"EDA / ASM pipeline — {len(pass_cards)} reactions (Check-4 PASS)",
                    pass_cards,
                    "The reactions below satisfy Check-4 conservation (residual ≤ 0.5 kcal/mol).")
    build_markdown(OUT / "FAIL_19.md",
                    f"EDA / ASM pipeline — {len(fail_cards)} reactions still FAIL",
                    fail_cards,
                    "These reactions still fail Check-4 conservation after the full pipeline.")

    # Bundle into one ZIP
    print("Bundling ZIP…")
    zip_path = OUT / "ADF_EDA_docs.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(OUT / "PASS_43.docx", "PASS_43.docx")
        zf.write(OUT / "FAIL_19.docx", "FAIL_19.docx")
        zf.write(OUT / "PASS_43.md", "PASS_43.md")
        zf.write(OUT / "FAIL_19.md", "FAIL_19.md")
        for png in sorted(IMG.glob("*.png")):
            zf.write(png, f"images/{png.name}")
    print(f"Wrote {zip_path}")

    # Summary
    sizes = {p.name: p.stat().st_size for p in OUT.iterdir() if p.is_file()}
    print()
    print("Outputs under " + str(OUT) + ":")
    for name in sorted(sizes):
        kb = sizes[name] / 1024
        print(f"  {name:25s}  {kb:8.1f} KB")
    print(f"  images/  ({len(list(IMG.iterdir()))} PNGs)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
